"""Helper module for AMIS automation and document manipulation.

Luồng làm việc:
- Đăng nhập -> mở trang chi tiết -> bấm More (ba chấm) -> 'In mẫu thiết lập'
- Chờ popup #popupexecution -> tick mẫu thứ 3 -> 'Tải xuống mẫu in đã chọn'
"""

import os
import time
import re
import requests
from typing import List, Tuple, Optional

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, ElementClickInterceptedException

from docx import Document
from docx.shared import Inches

# ===================== Selenium base =====================

def _make_driver(download_dir: str, headless: bool = True) -> webdriver.Chrome:
    os.makedirs(download_dir, exist_ok=True)
    opts = Options()
    if headless:
        opts.add_argument("--headless=new")
    opts.add_argument("--no-sandbox")
    opts.add_argument("--disable-dev-shm-usage")
    opts.add_argument("--disable-gpu")
    opts.add_argument("--window-size=1440,900")
    prefs = {
        "download.default_directory": download_dir,
        "download.prompt_for_download": False,
        "profile.default_content_settings.popups": 0,
        "safebrowsing.enabled": True,
    }
    opts.add_experimental_option("prefs", prefs)
    return webdriver.Chrome(options=opts)

# ===================== Utils =====================

def _log(msg: str):
    print(f"[AMIS] {msg}")

def _wait_css(driver, css: str, timeout: int = 15):
    return WebDriverWait(driver, timeout).until(
        EC.presence_of_element_located((By.CSS_SELECTOR, css))
    )

def _visible_and_clickable(driver, css: str, timeout: int = 15):
    return WebDriverWait(driver, timeout).until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, css))
    )

def _scroll_into_view_click(driver, el):
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
    try:
        el.click()
    except Exception:
        driver.execute_script("arguments[0].click();", el)

def _hover_then_click(driver, el):
    try:
        ActionChains(driver).move_to_element(el).pause(0.05).click(el).perform()
    except Exception:
        _scroll_into_view_click(driver, el)

def _try_click(driver, css: str, timeout: int = 8) -> bool:
    try:
        el = _visible_and_clickable(driver, css, timeout)
        _hover_then_click(driver, el)
        return True
    except Exception:
        try:
            el = driver.find_element(By.CSS_SELECTOR, css)
            _hover_then_click(driver, el)
            return True
        except Exception:
            return False

def _click_xpath(driver, xp: str, timeout: int = 15):
    el = WebDriverWait(driver, timeout).until(EC.element_to_be_clickable((By.XPATH, xp)))
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
    try:
        el.click()
    except Exception:
        driver.execute_script("arguments[0].click();", el)

def _try_click_xpath(driver, xp: str, timeout: int = 8) -> bool:
    try:
        el = WebDriverWait(driver, timeout).until(EC.element_to_be_clickable((By.XPATH, xp)))
        _hover_then_click(driver, el)
        return True
    except Exception:
        try:
            el = driver.find_element(By.XPATH, xp)
            _hover_then_click(driver, el)
            return True
        except Exception:
            return False

def _dump_debug(driver, out_dir: str, tag: str) -> None:
    try:
        os.makedirs(out_dir, exist_ok=True)
        driver.save_screenshot(os.path.join(out_dir, f"debug_{tag}.png"))
        with open(os.path.join(out_dir, f"debug_{tag}.html"), "w", encoding="utf-8") as f:
            f.write(driver.page_source)
        # dump thêm các iframe hiện có
        frames = driver.find_elements(By.CSS_SELECTOR, "iframe")
        with open(os.path.join(out_dir, f"debug_{tag}_iframes.txt"), "w", encoding="utf-8") as f:
            for i, fr in enumerate(frames):
                f.write(f"{i}: id={fr.get_attribute('id')} name={fr.get_attribute('name')} src={fr.get_attribute('src')}\n")
    except Exception:
        pass

# ===================== Frame helpers =====================

def _switch_into_notification_detail(driver, timeout: int = 15):
    driver.switch_to.default_content()
    WebDriverWait(driver, timeout).until(
        EC.frame_to_be_available_and_switch_to_it((By.CSS_SELECTOR, "iframe#notification-detail"))
    )

def _detail_ui_present_in_top_dom(driver) -> bool:
    """
    Trả về True nếu thấy dấu hiệu UI 'trang chi tiết' ngay ở DOM gốc (không trong iframe).
    """
    try:
        driver.switch_to.default_content()
        return bool(driver.find_elements(By.CSS_SELECTOR, "#top_nav, .wrap-icon-more.more-title-execution"))
    except Exception:
        return False

def _try_switch_into_any_detail_like_iframe(driver) -> bool:
    """
    Trong một số phiên bản, trang chi tiết KHÔNG nằm trong iframe.
    1) Nếu thấy UI chi tiết ở DOM gốc -> trả về True (không chuyển vào iframe).
    2) Nếu không thấy -> duyệt các iframe và chọn iframe chứa thanh top_nav / nút more.
    """
    # 1) Thử ở DOM gốc
    if _detail_ui_present_in_top_dom(driver):
        return True

    # 2) Nếu không có ở DOM gốc, duyệt các iframe
    driver.switch_to.default_content()
    frames = driver.find_elements(By.CSS_SELECTOR, "iframe")
    for fr in frames:
        try:
            driver.switch_to.default_content()
            driver.switch_to.frame(fr)
            if driver.find_elements(By.CSS_SELECTOR, "#top_nav, .wrap-icon-more.more-title-execution"):
                return True
        except Exception:
            continue
    driver.switch_to.default_content()
    return False

def _wait_popupexecution_anywhere(driver, timeout: int = 20) -> bool:
    end = time.time() + timeout
    while time.time() < end:
        # top
        try:
            driver.switch_to.default_content()
            if driver.find_elements(By.CSS_SELECTOR, "#popupexecution"):
                return True
        except Exception:
            pass
        # iframe chính
        try:
            if _try_switch_into_any_detail_like_iframe(driver):
                if driver.find_elements(By.CSS_SELECTOR, "#popupexecution"):
                    return True
        except Exception:
            pass
        time.sleep(0.2)
    return False

# ===================== Selectors =====================

# CSS fallback cũ
MORE_BTN_STRICT = (
    "div.nav.flex.items-center.offset-title-information > "
    "div.d-flex.content-user > "
    "div.d-flex.justify-flexend.wrap-icon-more.m-t-14.more-title-execution > button > div > i"
)
MORE_BTN_RELAX  = "div.d-flex.justify-flexend.wrap-icon-more.more-title-execution > button, .wrap-icon-more.more-title-execution > button"

# Fallback cuối cùng theo selector rất cụ thể
MORE_BTN_USER = (
    "#popupexecution > div.ms-popup.flex.flex-col > div.ms-popup--content-header > div > "
    "div:nth-child(2) > div > div > div.h-100.w-100.p-t-0.p-b-0.flex.flex-col > "
    "div.nav.flex.items-center.offset-title-information > div.d-flex.content-user > "
    "div.d-flex.justify-flexend.wrap-icon-more.m-t-14.more-title-execution > button > div > i"
)

IN_MAU_STRICT   = ("body > div.dx-overlay-wrapper.dx-popup-wrapper.dx-popover-wrapper.popover-action-process."
                   "dx-popover-without-title.dx-position-bottom > div > div.dx-popup-content > div > div:nth-child(2)")
IN_MAU_RELAX    = "div.dx-popover-wrapper.popover-action-process .dx-popup-content > div > div:nth-child(2)"

CHECKBOX_MAU3   = (
    "#popupexecution > div.ms-popup.flex.flex-col > div.ms-popup--content-header > div:nth-child(2) > "
    "div > div > div > div.dx-popup-content > div > div > div.dx-scrollable-wrapper > "
    "div > div.dx-scrollable-content > div.dx-scrollview-content > div > div > "
    "div.list-template.h-full.pos-relative > div.h-full.p-16 > div:nth-child(3) > label > "
    "span.icon-square-check-primary.checkmark"
)

DOWNLOAD_BLUE   = (
    "#popupexecution > div.ms-popup.flex.flex-col > div.ms-popup--content-header > div:nth-child(2) > "
    "div > div > div > div.dx-popup-content > div > div > div.dx-scrollable-wrapper > "
    "div > div.dx-scrollable-content > div.dx-scrollview-content > div > div > "
    "div.list-template.h-full.pos-relative > div.h-full.p-16 > div.flex.items-center.justify-between.m-b-8 > "
    "div > div.text-blue"
)

POPOVER_WRAPPER = "body div.dx-popover-wrapper.popover-action-process, body div.dx-popup-wrapper.popover-action-process"
POPUPEXECUTION  = "#popupexecution"

# ===== XPATH tuyệt đối (ưu tiên thử trước) =====
XPATH_MORE_BTN       = "/html/body/div[1]/div[2]/div[2]/div/div[2]/div/div/div[1]/div[1]/div[2]/div[2]/button/div/i"
XPATH_IN_MAU         = "/html/body/div[12]/div/div[2]/div/div[2]"
XPATH_CHECKBOX_MAU3  = "/html/body/div[1]/div[2]/div[2]/div[2]/div/div/div/div[2]/div/div/div[1]/div/div[1]/div[2]/div/div/div[1]/div[2]/div[3]/label/span[1]"
XPATH_DOWNLOAD_BLUE  = "/html/body/div[1]/div[2]/div[2]/div[2]/div/div/div/div[2]/div/div/div[1]/div/div[1]/div[2]/div/div/div[1]/div[2]/div[1]/div/div[2]"

# ===================== Click “In mẫu…” ổn định theo TEXT =====================

def _click_in_mau_anywhere(driver, timeout: int = 12) -> bool:
    """
    Click item 'In mẫu thiết lập' / 'Xem trước mẫu in' / 'In mẫu' / 'Mẫu in' / 'Xem trước'
    trong các popover DevExtreme, thử ở top và trong iframe; không phụ thuộc index.
    """
    WANT_TEXTS = ["In mẫu thiết lập", "Xem trước mẫu in", "In mẫu", "Mẫu in", "Xem trước"]

    def _click_candidate(el) -> bool:
        try:
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            try:
                el.click()
            except (ElementClickInterceptedException, Exception):
                # fallback click bằng JS theo toạ độ ở giữa phần tử
                driver.execute_script("""
const el = arguments[0];
const r = el.getBoundingClientRect();
const x = r.left + r.width/2, y = r.top + r.height/2;
const ev = (type)=>{ const e=new MouseEvent(type,{bubbles:true,cancelable:true,clientX:x,clientY:y}); el.dispatchEvent(e); };
ev('pointerdown'); ev('mousedown'); ev('mouseup'); ev('click');
""", el)
            return True
        except Exception:
            return False

    def _try_in_context(switch_into_iframe: bool) -> bool:
        try:
            driver.switch_to.default_content()
            if switch_into_iframe:
                if not _try_switch_into_any_detail_like_iframe(driver):
                    return False
        except Exception:
            return False

        # chờ popover xuất hiện
        end_pop = time.time() + 3
        while time.time() < end_pop:
            if driver.find_elements(By.CSS_SELECTOR, POPOVER_WRAPPER):
                break
            time.sleep(0.1)

        # thử tìm theo text trong wrapper
        wrappers = driver.find_elements(By.CSS_SELECTOR, POPOVER_WRAPPER)
        for w in wrappers:
            # duyệt children để tránh click nhầm wrapper
            cands = w.find_elements(By.CSS_SELECTOR, "*")
            for el in cands:
                try:
                    txt = (el.text or "").strip()
                    if not txt:
                        continue
                    for want in WANT_TEXTS:
                        if want in txt:
                            if _click_candidate(el):
                                return True
                except Exception:
                    continue

        # Nếu wrapper không có text (một số case), thử tìm theo role/button
        # và so sánh text gần kề
        try:
            buttons = driver.find_elements(By.CSS_SELECTOR, "div[role='button'], .dx-item, .dx-button, .dx-menu-item")
            for el in buttons:
                t = (el.text or "").strip()
                for want in WANT_TEXTS:
                    if want in t and _click_candidate(el):
                        return True
        except Exception:
            pass

        return False

    end = time.time() + timeout
    while time.time() < end:
        if _try_in_context(False) or _try_in_context(True):
            return True
        time.sleep(0.2)
    return False

# ===================== Open popover & click “In mẫu thiết lập” =====================

def _open_print_preview_via_popover(driver, download_dir: str) -> None:
    """
    - Click nút More TRONG iframe (ưu tiên XPath bạn đưa; fallback CSS)
    - Quay về default_content; click item “In mẫu…” bằng TEXT (ổn định)
    - Đợi #popupexecution xuất hiện ở top hoặc trong iframe
    """
    # Nếu popup đã có thì thôi
    if _wait_popupexecution_anywhere(driver, timeout=1):
        _log("Popup #popupexecution đã mở sẵn.")
        return

    # (A) Chuẩn bị ngữ cảnh (iframe hoặc DOM gốc)
    in_detail_context = _try_switch_into_any_detail_like_iframe(driver)
    if not in_detail_context:
        # Thử iframe chuẩn nếu có
        try:
            _switch_into_notification_detail(driver, timeout=20)
            in_detail_context = True
            _log("Đã switch vào iframe notification-detail.")
        except Exception:
            # Không ép raise nữa; có thể trang chi tiết đang ở DOM gốc
            driver.switch_to.default_content()
            in_detail_context = _detail_ui_present_in_top_dom(driver)
            if not in_detail_context:
                _dump_debug(driver, download_dir, "cannot_find_detail_iframe_but_try_top_dom")
                # Không raise tại đây; để phần click 'More' tự xử lý fallback

    try:
        # ưu tiên XPath
        if not _try_click_xpath(driver, XPATH_MORE_BTN, timeout=8):
            # fallback CSS
            if not (_try_click(driver, MORE_BTN_STRICT, timeout=6) or _try_click(driver, MORE_BTN_RELAX, timeout=6)):
                # Fallback cuối theo selector rất cụ thể
                if not _try_click(driver, MORE_BTN_USER, timeout=4):
                    _dump_debug(driver, download_dir, "cannot_click_more_any")
                    raise TimeoutException("Không click được nút 3 chấm (More).")
        _log("Đã click nút 3 chấm.")
    except Exception as e:
        _dump_debug(driver, download_dir, "cannot_click_more_exception")
        raise TimeoutException(f"Không click được nút 3 chấm: {e}")

    # (B) Click “In mẫu…” KHÔNG phụ thuộc div[12]
    driver.switch_to.default_content()
    # đợi popover hiện ra
    end_wait = time.time() + 6
    while time.time() < end_wait:
        if driver.find_elements(By.CSS_SELECTOR, POPOVER_WRAPPER):
            break
        time.sleep(0.1)

    # thử 1: XPath tuyệt đối (nếu DOM đúng)
    clicked = _try_click_xpath(driver, XPATH_IN_MAU, timeout=2)
    if clicked:
        _log("Click 'In mẫu...' theo XPath tuyệt đối (có index).")
    else:
        # thử 2: CSS cũ
        clicked = _try_click(driver, IN_MAU_STRICT, timeout=3) or _try_click(driver, IN_MAU_RELAX, timeout=3)
        if clicked:
            _log("Click 'In mẫu...' theo CSS fallback.")
        else:
            # thử 3: theo TEXT (ổn định)
            clicked = _click_in_mau_anywhere(driver, timeout=10)
            if clicked:
                _log("Click 'In mẫu...' bằng TEXT (ổn định).")

    if not clicked:
        _dump_debug(driver, download_dir, "fail_click_in_mau_all_methods")
        raise TimeoutException("Không mở được 'In mẫu thiết lập' (tất cả phương án click đều fail).")

    # (C) Đợi popup xuất hiện (ở top hoặc trong iframe)
    if not _wait_popupexecution_anywhere(driver, timeout=25):
        _dump_debug(driver, download_dir, "no_popupexecution_after_click_in_mau")
        raise TimeoutException("Không mở được 'In mẫu thiết lập' (không thấy #popupexecution).")

# ===================== Chọn template & tải xuống =====================

def _choose_template_and_download(driver, download_dir: str) -> str:
    """
    Chọn mẫu thứ 3 và tải .docx. Ưu tiên XPath bạn cung cấp, có fallback CSS và tìm theo text trong #popupexecution.
    """
    if not _wait_popupexecution_anywhere(driver, timeout=25):
        _dump_debug(driver, download_dir, "no_popup_when_choose_template")
        raise TimeoutException("Popup 'Xem trước mẫu in' không xuất hiện.")

    # Control thường nằm trong iframe trang chi tiết, nhưng có thể ở DOM gốc
    if not _try_switch_into_any_detail_like_iframe(driver):
        try:
            _switch_into_notification_detail(driver, timeout=15)
        except Exception:
            driver.switch_to.default_content()
            if not _detail_ui_present_in_top_dom(driver):
                _dump_debug(driver, download_dir, "cannot_switch_iframe_for_popup")
                raise TimeoutException("Không vào được iframe để thao tác popup.")

    # Tick mẫu thứ 3
    try:
        if not _try_click_xpath(driver, XPATH_CHECKBOX_MAU3, timeout=10):
            if not _try_click(driver, CHECKBOX_MAU3, timeout=8):
                _dump_debug(driver, download_dir, "cannot_tick_template_3")
                raise TimeoutException("Không tick được mẫu thứ 3.")
        _log("Đã tick mẫu thứ 3.")
    except Exception:
        _dump_debug(driver, download_dir, "tick_template_3_exception")
        raise

    time.sleep(0.3)

    # Bấm “Tải mẫu in”
    try:
        if not _try_click_xpath(driver, XPATH_DOWNLOAD_BLUE, timeout=10):
            if not _try_click(driver, DOWNLOAD_BLUE, timeout=8):
                # fallback cuối: tìm theo text trong #popupexecution
                try:
                    driver.execute_script("""
const root=document.querySelector('#popupexecution'); if(!root) return;
function vis(el){if(!el)return false; const s=getComputedStyle(el);
  if(s.display==='none'||s.visibility==='hidden') return false;
  const r=el.getBoundingClientRect(); return r.width>0 && r.height>0; }
const wants=['Tải mẫu in','Tải xuống','Tải về','Download'];
const all=root.querySelectorAll('*');
for(const el of all){ if(!vis(el)) continue; const t=(el.innerText||'').trim(); if(!t) continue;
  for(const w of wants){ if(t.includes(w)){ try{ el.scrollIntoView({block:'center'}); el.click(); return; }catch(e){} } } }
""")
                except Exception:
                    pass
        _log("Đã bấm 'Tải mẫu in'.")
    except Exception:
        _dump_debug(driver, download_dir, "click_download_exception")
        raise

    return _wait_for_docx(download_dir, timeout=120)

# ===================== MAIN =====================

def run_automation(
    username: str,
    password: str,
    download_dir: str,
    headless: bool = True,
    record_id: Optional[str] = None,
    execution_id: Optional[str] = None,
    status: int = 1,
    company_code: str = "RH7VZQAQ",
) -> Tuple[str, List[str]]:
    if not execution_id and record_id:
        execution_id = str(record_id)
    if not execution_id:
        raise ValueError("Cần truyền execution_id (hoặc record_id sẽ được coi là execution_id).")

    driver = _make_driver(download_dir, headless=headless)
    wait = WebDriverWait(driver, 90)

    try:
        # 1) Login
        driver.get("https://amisapp.misa.vn/")
        user_el = wait.until(
            EC.presence_of_element_located(
                (By.CSS_SELECTOR, "#box-login-right .login-form-inputs .username-wrap input")
            )
        )
        user_el.send_keys(username)
        driver.find_element(By.CSS_SELECTOR, "#box-login-right .login-form-inputs .pass-wrap input").send_keys(password)
        driver.find_element(By.CSS_SELECTOR, "#box-login-right .login-form-btn-container button").click()
        wait.until(EC.url_contains("amisapp.misa.vn"))
        time.sleep(0.6)

        # Best-effort đóng các pop-up chào mừng
        for by, sel in [
            (By.XPATH, "//button[contains(.,'Bỏ qua')]"),
            (By.XPATH, "//button[contains(.,'Tiếp tục làm việc')]"),
            (By.XPATH, "//button[contains(.,'Đóng')]"),
            (By.CSS_SELECTOR, "[aria-label='Close'],[data-dismiss],.close"),
        ]:
            try:
                driver.find_element(by, sel).click()
                time.sleep(0.05)
            except Exception:
                pass

        # 2) Vào trang chi tiết
        detail_url = (
            "https://amisapp.misa.vn/process/execute/1"
            f"?ID={execution_id}&type=1&status={status}&appCode=AMISProcess&companyCode={company_code}"
        )
        driver.get(detail_url)
        try:
            WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.CSS_SELECTOR, "#top_nav")))
        except Exception:
            pass
        time.sleep(0.3)

        # 3) Mở popover và click 'In mẫu thiết lập'
        _open_print_preview_via_popover(driver, download_dir)

        # 4) Chọn mẫu & tải docx
        template_path = _choose_template_and_download(driver, download_dir)

        # 5) Tải ảnh liên quan (best-effort)
        images = _download_images_from_detail(driver, download_dir)

        return template_path, images
    finally:
        driver.quit()

# ===================== Helpers: files & docx =====================

def _wait_for_docx(folder: str, timeout: int = 120) -> str:
    for _ in range(timeout):
        for f in os.listdir(folder):
            if f.lower().endswith(".docx"):
                return os.path.join(folder, f)
        time.sleep(1)
    raise FileNotFoundError("Không thấy file .docx sau khi tải")

def _download_images_from_detail(driver, download_dir: str) -> List[str]:
    images: List[str] = []
    for i, t in enumerate(driver.find_elements(By.CSS_SELECTOR, "img")[:8], start=1):
        try:
            src = t.get_attribute("src")
            if src and src.startswith("http"):
                r = requests.get(src, timeout=15)
                cd = r.headers.get("Content-Disposition", "")
                m = re.search(r'filename="?([^\"]+)"?', cd) if cd else None
                name = m.group(1) if m else f"image_{i}.jpg"
                p = os.path.join(download_dir, name)
                with open(p, "wb") as f:
                    f.write(r.content)
                images.append(p)
        except Exception:
            pass
    return images

# ===================== Word processing =====================

def fill_document(template_path: str, images: List[str], signature_path: str, output_path: str) -> None:
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)
    slot_map = {
        "Thông tin rao bán/sổ đỏ": images[0:1],
        "Mặt trước tài sản": images[1:2],
        "Tổng thể tài sản": images[2:3],
        "Đường phía trước tài sản": images[3:5],
        "Ảnh khác": images[5:7],
    }

    def _table_has_phu_luc(tbl):
        text = "\n".join(cell.text for row in tbl.rows for cell in row.cells)
        return ("Phụ lục" in text) and ("Ảnh TSSS" in text)

    target_table = None
    for tbl in doc.tables:
        if _table_has_phu_luc(tbl):
            target_table = tbl
            break

    if target_table:
        for row in target_table.rows:
            for ci, cell in enumerate(row.cells):
                label = cell.text.strip()
                if label in slot_map and ci + 1 < len(row.cells):
                    dest = row.cells[ci + 1]
                    for p in dest.paragraphs:
                        if p.text:
                            p.text = ""
                    for pth in slot_map[label]:
                        if os.path.exists(pth):
                            dest.paragraphs[0].add_run().add_picture(pth, width=Inches(2.2))
                            dest.add_paragraph("")
    else:
        doc.add_heading("Phụ lục Ảnh TSSS", level=2)
        for pth in images:
            if os.path.exists(pth):
                doc.add_picture(pth, width=Inches(3))
                doc.add_paragraph("")

    doc.add_page_break()
    doc.add_heading("Chữ ký cán bộ khảo sát", level=2)
    if signature_path and os.path.exists(signature_path):
        try:
            doc.add_picture(signature_path, width=Inches(2))
        except Exception as e:
            doc.add_paragraph(f"[Không chèn được chữ ký: {e}]")
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
